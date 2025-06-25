import os
import pypdf
from docx import Document
from pdf2docx import Converter as Pdf2DocxConverter
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import threading
import time

# --- Configure Tesseract Path for macOS ---
# Check if tesseract is available in PATH, otherwise set the path manually
try:
    import subprocess
    result = subprocess.run(['which', 'tesseract'], capture_output=True, text=True)
    if result.returncode == 0:
        tesseract_path = result.stdout.strip()
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        print(f"Tesseract найден: {tesseract_path}")
    else:
        # Fallback path for macOS with Homebrew
        pytesseract.pytesseract.tesseract_cmd = '/opt/homebrew/bin/tesseract'
        print("Используется стандартный путь для macOS Homebrew")
except Exception as e:
    print(f"Ошибка при настройке Tesseract: {e}")
    # Fallback path
    pytesseract.pytesseract.tesseract_cmd = '/opt/homebrew/bin/tesseract'
# -----------------------------------------------------------------------------

def extract_text_from_pdf_pypdf(pdf_path):
    """Attempts to extract text directly from a PDF using pypdf."""
    try:
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = pypdf.PdfReader(pdf_file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text() or ""
        return text
    except Exception as e:
        raise Exception(f"Ошибка при извлечении текста с помощью pypdf: {e}")

def ocr_pdf_to_txt(pdf_path, output_folder, lang='rus+eng'):
    """Performs OCR on a PDF file and saves the text to a TXT file."""
    try:
        # First, test if tesseract is working
        try:
            test_result = pytesseract.get_tesseract_version()
            print(f"Tesseract версия: {test_result}")
        except Exception as e:
            raise Exception(f"Tesseract не работает: {e}")
        
        # Check if language is available
        try:
            available_langs = pytesseract.get_languages()
            if 'rus' not in available_langs:
                print("Предупреждение: русский язык не найден, используем английский")
                lang = 'eng'
        except Exception as e:
            print(f"Не удалось проверить языки: {e}")
            lang = 'eng'
        
        # Convert PDF to images with higher DPI for better OCR
        try:
            images = convert_from_path(pdf_path, dpi=300)
            print(f"PDF конвертирован в {len(images)} изображений")
        except Exception as e:
            raise Exception(f"Ошибка при конвертации PDF в изображения: {e}")
        
        full_text = []
        for i, image in enumerate(images):
            print(f"Обрабатывается страница {i+1}/{len(images)}")
            
            # Convert PIL image to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Try different OCR configurations for better results
            ocr_configs = [
                r'--oem 3 --psm 6',
                r'--oem 3 --psm 3',
                r'--oem 3 --psm 1'
            ]
            
            best_text = ""
            best_confidence = 0
            
            for config in ocr_configs:
                try:
                    data = pytesseract.image_to_data(image, lang=lang, output_type=pytesseract.Output.DICT, config=config)
                    
                    confidences = [conf for conf in data['conf'] if conf > 0]
                    avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                    
                    lines = {}
                    for j in range(len(data['text'])):
                        word = data['text'][j]
                        left = data['left'][j]
                        top = data['top'][j]
                        conf = data['conf'][j]
                        
                        if word.strip() and conf > 20:
                            line_key = top // 15
                            
                            if line_key not in lines:
                                lines[line_key] = []
                            
                            lines[line_key].append((left, word, conf))
                    
                    sorted_line_keys = sorted(lines.keys())
                    page_text = []
                    for line_key in sorted_line_keys:
                        sorted_words = sorted(lines[line_key])
                        
                        current_line = []
                        prev_right = 0
                        for left, word, conf in sorted_words:
                            if left > prev_right:
                                spaces = max(1, (left - prev_right) // 8)
                                current_line.append(' ' * spaces)
                            current_line.append(word)
                            prev_right = left + len(word) * 8
                        
                        line_text = ''.join(current_line).strip()
                        if line_text:
                            page_text.append(line_text)
                    
                    current_text = '\n'.join(page_text)
                    
                    if avg_confidence > best_confidence or (avg_confidence == best_confidence and len(current_text) > len(best_text)):
                        best_text = current_text
                        best_confidence = avg_confidence
                        
                except Exception as e:
                    print(f"Ошибка с конфигурацией {config}: {e}")
                    continue
            
            if best_text:
                full_text.append(best_text)
                print(f"Страница {i+1}: найдено {len(best_text)} символов (уверенность: {best_confidence:.1f}%)")
            else:
                print(f"Страница {i+1}: текст не найден")
                full_text.append(f"[Страница {i+1}: текст не распознан]")
            
            if i < len(images) - 1:
                full_text.append("\n--- Страница {} ---\n".format(i + 2))

        output_text = "\n".join(full_text)

        os.makedirs(output_folder, exist_ok=True)
        filename_without_ext = os.path.splitext(os.path.basename(pdf_path))[0]
        txt_output_path = os.path.join(output_folder, f"{filename_without_ext}.txt")
        
        with open(txt_output_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(output_text)
        
        return True, f"Успешно конвертировано (OCR): {os.path.basename(pdf_path)}"
    except Exception as e:
        raise Exception(f"Ошибка при конвертации с помощью OCR: {e}")

def convert_pdf_to_txt_direct(pdf_path, output_folder):
    """Converts a PDF file directly to a TXT file using pypdf."""
    try:
        text = extract_text_from_pdf_pypdf(pdf_path)
        
        os.makedirs(output_folder, exist_ok=True)
        filename_without_ext = os.path.splitext(os.path.basename(pdf_path))[0]
        txt_output_path = os.path.join(output_folder, f"{filename_without_ext}.txt")
        
        with open(txt_output_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(text)
        return True, f"Успешно конвертировано (прямо): {os.path.basename(pdf_path)}"
    except Exception as e:
        raise Exception(f"Ошибка при прямой конвертации в TXT: {e}")

def convert_pdf_to_docx_then_txt(pdf_path, output_folder):
    """Converts a PDF file to DOCX and then extracts text from the DOCX to TXT."""
    try:
        os.makedirs(output_folder, exist_ok=True)
        
        filename_without_ext = os.path.splitext(os.path.basename(pdf_path))[0]
        docx_temp_path = os.path.join(output_folder, f"temp_{filename_without_ext}.docx")
        txt_output_path = os.path.join(output_folder, f"{filename_without_ext}.txt")

        cv = Pdf2DocxConverter(pdf_path)
        cv.convert(docx_temp_path)
        cv.close()

        doc = Document(docx_temp_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)

        with open(txt_output_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(text)

        os.remove(docx_temp_path)
        
        return True, f"Успешно конвертировано (через DOCX): {os.path.basename(pdf_path)}"
    except Exception as e:
        raise Exception(f"Ошибка при конвертации через DOCX в TXT: {e}")

def convert_docx_to_txt(docx_path, output_folder):
    """Конвертирует DOCX файл в TXT с кодировкой utf-8."""
    try:
        from docx import Document
        import os
        os.makedirs(output_folder, exist_ok=True)
        filename_without_ext = os.path.splitext(os.path.basename(docx_path))[0]
        txt_output_path = os.path.join(output_folder, f"{filename_without_ext}.txt")
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)
        with open(txt_output_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(text)
        return True, f"Успешно конвертировано DOCX -> TXT: {os.path.basename(docx_path)}"
    except Exception as e:
        return False, f"Ошибка при конвертации DOCX -> TXT: {e}"

def select_files():
    file_paths = filedialog.askopenfilenames(
        title="Выберите PDF файлы для конвертации",
        filetypes=[("PDF files", "*.pdf")]
    )
    return file_paths

def select_output_folder():
    folder_path = filedialog.askdirectory(
        title="Выберите папку для сохранения TXT файлов"
    )
    return folder_path

class ConversionProgress:
    def __init__(self, root):
        self.root = root
        self.progress_window = None
        self.progress_bar = None
        self.status_label = None
        self.current_file_label = None
        self.results_text = None
        self.close_button = None
    
    def show_progress_window(self, total_files):
        self.progress_window = tk.Toplevel(self.root)
        self.progress_window.title("Прогресс конвертации")
        self.progress_window.geometry("600x430")
        self.progress_window.transient(self.root)
        self.progress_window.grab_set()
        
        # Progress bar
        self.progress_bar = ttk.Progressbar(self.progress_window, length=500, mode='determinate')
        self.progress_bar.pack(pady=10)
        
        # Status labels
        self.status_label = tk.Label(self.progress_window, text=f"Обработано: 0 из {total_files}")
        self.status_label.pack(pady=5)
        
        self.current_file_label = tk.Label(self.progress_window, text="", wraplength=550)
        self.current_file_label.pack(pady=5)
        
        # Results text area
        results_frame = tk.Frame(self.progress_window)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        tk.Label(results_frame, text="Результаты:").pack(anchor=tk.W)
        
        self.results_text = tk.Text(results_frame, height=15, width=70)
        scrollbar = tk.Scrollbar(results_frame, orient=tk.VERTICAL, command=self.results_text.yview)
        self.results_text.configure(yscrollcommand=scrollbar.set)
        
        self.results_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Close button (disabled initially)
        self.close_button = tk.Button(self.progress_window, text="Закрыть", state=tk.DISABLED, command=self.close)
        self.close_button.pack(pady=10)
    
    def update_progress(self, current, total, current_file=""):
        if self.progress_bar:
            progress = (current / total) * 100
            self.progress_bar['value'] = progress
            self.status_label.config(text=f"Обработано: {current} из {total}")
            self.current_file_label.config(text=f"Текущий файл: {os.path.basename(current_file) if current_file else ''}")
            self.progress_window.update()
    
    def add_result(self, message, is_error=False):
        if self.results_text:
            timestamp = time.strftime("%H:%M:%S")
            color = "red" if is_error else "green"
            self.results_text.insert(tk.END, f"[{timestamp}] {message}\n")
            self.results_text.see(tk.END)
            self.progress_window.update()
    
    def enable_close(self):
        if self.close_button:
            self.close_button.config(state=tk.NORMAL)
    
    def close(self):
        if self.progress_window:
            self.progress_window.destroy()

def start_conversion(conversion_method):
    pdf_files = select_files()
    if not pdf_files:
        messagebox.showinfo("Информация", "Файлы не выбраны.")
        return

    output_folder = select_output_folder()
    if not output_folder:
        messagebox.showinfo("Информация", "Папка для сохранения не выбрана.")
        return

    # Create progress window
    progress = ConversionProgress(root)
    progress.show_progress_window(len(pdf_files))
    
    def conversion_worker():
        successful_conversions = []
        failed_conversions = []
        
        for i, pdf_file_path in enumerate(pdf_files):
            try:
                progress.update_progress(i, len(pdf_files), pdf_file_path)
                
                success, message = False, "Ошибка: неверный метод конвертации."
                if conversion_method == 'ocr':
                    success, message = ocr_pdf_to_txt(pdf_file_path, output_folder)
                elif conversion_method == 'direct_txt':
                    success, message = convert_pdf_to_txt_direct(pdf_file_path, output_folder)
                elif conversion_method == 'docx_then_txt':
                    success, message = convert_pdf_to_docx_then_txt(pdf_file_path, output_folder)
                
                if success:
                    successful_conversions.append(os.path.basename(pdf_file_path))
                    progress.add_result(f"✅ {message}")
                else:
                    failed_conversions.append(os.path.basename(pdf_file_path))
                    progress.add_result(f"❌ {message}", is_error=True)
                    
            except Exception as e:
                failed_conversions.append(os.path.basename(pdf_file_path))
                error_msg = f"❌ Ошибка при конвертации {os.path.basename(pdf_file_path)}: {str(e)}"
                progress.add_result(error_msg, is_error=True)
                print(f"Ошибка при конвертации {pdf_file_path}: {e}")
        
        # Final update
        progress.update_progress(len(pdf_files), len(pdf_files))
        
        # Show final results
        final_message = f"Конвертация завершена!\n\n"
        final_message += f"✅ Успешно конвертировано: {len(successful_conversions)}\n"
        final_message += f"❌ Ошибок: {len(failed_conversions)}\n\n"
        
        if failed_conversions:
            final_message += "Список файлов с ошибками (копируйте для поиска):\n"
            for failed_file in failed_conversions:
                final_message += f"{failed_file}\n"
        
        # Save error report
        if failed_conversions:
            error_report_path = os.path.join(output_folder, "error_report.txt")
            with open(error_report_path, 'w', encoding='utf-8') as f:
                f.write("Отчет об ошибках конвертации\n")
                f.write("=" * 40 + "\n\n")
                f.write(f"Дата: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Всего файлов: {len(pdf_files)}\n")
                f.write(f"Успешно: {len(successful_conversions)}\n")
                f.write(f"Ошибок: {len(failed_conversions)}\n\n")
                f.write("Список файлов с ошибками (копируйте для поиска):\n")
                for failed_file in failed_conversions:
                    f.write(f"{failed_file}\n")
            
            final_message += f"\nОтчет об ошибках сохранен в: {os.path.basename(error_report_path)}"
        
        progress.add_result("\n" + "="*50)
        progress.add_result(final_message)
        
        # Show final message box
        root.after(1000, lambda: messagebox.showinfo("Конвертация завершена", final_message))
        
        # Enable close button (manual close only)
        progress.enable_close()
    
    # Start conversion in separate thread
    conversion_thread = threading.Thread(target=conversion_worker)
    conversion_thread.daemon = True
    conversion_thread.start()

def start_docx_to_txt_conversion():
    docx_files = filedialog.askopenfilenames(
        title="Выберите DOCX файлы для конвертации",
        filetypes=[("DOCX files", "*.docx")]
    )
    if not docx_files:
        messagebox.showinfo("Информация", "Файлы не выбраны.")
        return

    output_folder = select_output_folder()
    if not output_folder:
        messagebox.showinfo("Информация", "Папка для сохранения не выбрана.")
        return

    progress = ConversionProgress(root)
    progress.show_progress_window(len(docx_files))

    def conversion_worker():
        successful_conversions = []
        failed_conversions = []
        for i, docx_file_path in enumerate(docx_files):
            try:
                progress.update_progress(i, len(docx_files), docx_file_path)
                success, message = convert_docx_to_txt(docx_file_path, output_folder)
                if success:
                    successful_conversions.append(os.path.basename(docx_file_path))
                    progress.add_result(f"✅ {message}")
                else:
                    failed_conversions.append(os.path.basename(docx_file_path))
                    progress.add_result(f"❌ {message}", is_error=True)
            except Exception as e:
                failed_conversions.append(os.path.basename(docx_file_path))
                error_msg = f"❌ Ошибка при конвертации {os.path.basename(docx_file_path)}: {str(e)}"
                progress.add_result(error_msg, is_error=True)
                print(f"Ошибка при конвертации {docx_file_path}: {e}")
        progress.update_progress(len(docx_files), len(docx_files))
        final_message = f"Конвертация завершена!\n\n"
        final_message += f"✅ Успешно конвертировано: {len(successful_conversions)}\n"
        final_message += f"❌ Ошибок: {len(failed_conversions)}\n\n"
        if failed_conversions:
            final_message += "Список файлов с ошибками (копируйте для поиска):\n"
            for failed_file in failed_conversions:
                final_message += f"{failed_file}\n"
            error_report_path = os.path.join(output_folder, "error_report.txt")
            with open(error_report_path, 'w', encoding='utf-8') as f:
                f.write("Отчет об ошибках конвертации\n")
                f.write("=" * 40 + "\n\n")
                f.write(f"Дата: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Всего файлов: {len(docx_files)}\n")
                f.write(f"Успешно: {len(successful_conversions)}\n")
                f.write(f"Ошибок: {len(failed_conversions)}\n\n")
                f.write("Список файлов с ошибками (копируйте для поиска):\n")
                for failed_file in failed_conversions:
                    f.write(f"{failed_file}\n")
            final_message += f"\nОтчет об ошибках сохранен в: {os.path.basename(error_report_path)}"
        progress.add_result("\n" + "="*50)
        progress.add_result(final_message)
        root.after(1000, lambda: messagebox.showinfo("Конвертация завершена", final_message))
        progress.enable_close()
    conversion_thread = threading.Thread(target=conversion_worker)
    conversion_thread.daemon = True
    conversion_thread.start()

# GUI Setup
root = tk.Tk()
root.title("PDF в TXT конвертер (Улучшенный)")
root.geometry("500x300")

# Instruction Label
instruction_label = tk.Label(root, text="1. Выберите PDF файлы.\n2. Выберите папку для сохранения.\n3. Выберите метод конвертации.\n\nФайлы с ошибками будут пропущены и добавлены в отчет.", justify=tk.LEFT)
instruction_label.pack(pady=10)

# Buttons for conversion methods
ocr_button = tk.Button(root, text="OCR сканирование (для сканированных PDF и таблиц)", command=lambda: start_conversion('ocr'))
ocr_button.pack(pady=5, fill=tk.X, padx=20)

direct_txt_button = tk.Button(root, text="PDF -> TXT (прямая конвертация, для простых PDF)", command=lambda: start_conversion('direct_txt'))
direct_txt_button.pack(pady=5, fill=tk.X, padx=20)

docx_txt_button = tk.Button(root, text="PDF -> DOCX -> TXT (для проблемных PDF)", command=lambda: start_conversion('docx_then_txt'))
docx_txt_button.pack(pady=5, fill=tk.X, padx=20)

# Добавляю кнопку для DOCX -> TXT
docx2txt_button = tk.Button(root, text="DOCX -> TXT (конвертация DOCX в TXT)", command=start_docx_to_txt_conversion)
docx2txt_button.pack(pady=5, fill=tk.X, padx=20)

root.mainloop()